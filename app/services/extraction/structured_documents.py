from __future__ import annotations

import logging
import re
import unicodedata
from datetime import date, datetime
from pathlib import Path
from threading import Lock
from typing import Iterable

import pandas as pd
from openpyxl import load_workbook

from app.services.extraction.file_types import (
    MESSAGE_SOURCE_TYPES,
    SPREADSHEET_SOURCE_TYPES,
    STRUCTURED_TEXT_SOURCE_TYPES,
    WORD_SOURCE_TYPES,
)

logger = logging.getLogger(__name__)

_DOC_COM_LOCK = Lock()
SPREADSHEET_FIELD_LABELS = {
    "RAZON SOCIAL": "Razon Social",
    "RAZON SOCIAL PROVEEDOR": "Razon Social Proveedor",
    "NOMBRE PROVEEDOR": "Proveedor",
    "NOMBRE DEL PROVEEDOR": "Proveedor",
    "PROVEEDOR": "Proveedor",
    "EMISOR": "Emisor",
    "RFC": "RFC",
    "R F C": "RFC",
    "REGISTRO FEDERAL DE CONTRIBUYENTES": "RFC",
}
SPREADSHEET_LABEL_NOISE = {
    "CATEGORIA",
    "ELABORA",
    "NEGOCIO",
    "VIA",
    "FECHA",
    "TIPO DE NEGOCIO",
    "CODIGO DE BARRAS",
    "DESCRIPCION",
    "COSTO BRUTO",
    "COSTO NETO",
}
SPREADSHEET_NEARBY_OFFSETS = (
    (1, 0),
    (0, 1),
    (1, 1),
    (1, -1),
    (0, -1),
    (-1, 0),
    (-1, 1),
    (-1, -1),
    (2, 0),
    (0, 2),
    (2, 1),
    (1, 2),
)


def _clean_text(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ", timespec="seconds")
    if isinstance(value, date):
        return value.isoformat()

    text = str(value).replace("\x00", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def _normalize_spreadsheet_token(value: object) -> str:
    text = _clean_text(value).upper()
    text = unicodedata.normalize("NFD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = re.sub(r"[^A-Z0-9&]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _spreadsheet_field_label(value: object) -> str | None:
    normalized = _normalize_spreadsheet_token(value)
    if not normalized:
        return None
    if normalized in SPREADSHEET_FIELD_LABELS:
        return SPREADSHEET_FIELD_LABELS[normalized]
    return None


def _looks_like_spreadsheet_value(value: object) -> bool:
    cleaned = _clean_text(value)
    if not cleaned:
        return False
    if len(cleaned) < 4:
        return False

    normalized = _normalize_spreadsheet_token(cleaned)
    if not normalized:
        return False
    if normalized in SPREADSHEET_FIELD_LABELS or normalized in SPREADSHEET_LABEL_NOISE:
        return False
    if re.fullmatch(r"[\d\s\-/.,$%]+", cleaned):
        return False

    letters = [char for char in cleaned if char.isalpha()]
    return len(letters) >= 3


def _cell_value(rows: list[list[object]], row_index: int, col_index: int) -> object:
    if row_index < 0 or row_index >= len(rows):
        return None
    row = rows[row_index]
    if col_index < 0 or col_index >= len(row):
        return None
    return row[col_index]


def _spreadsheet_nearby_value(rows: list[list[object]], row_index: int, col_index: int) -> str | None:
    for row_offset, col_offset in SPREADSHEET_NEARBY_OFFSETS:
        candidate = _cell_value(rows, row_index + row_offset, col_index + col_offset)
        if _looks_like_spreadsheet_value(candidate):
            return _clean_text(candidate)
    return None


def _extract_spreadsheet_field_hints(rows: list[list[object]]) -> list[str]:
    hints: list[str] = []
    seen: set[tuple[str, str]] = set()

    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            label = _spreadsheet_field_label(value)
            if not label:
                continue

            nearby_value = _spreadsheet_nearby_value(rows, row_index, col_index)
            if not nearby_value:
                continue

            key = (label, nearby_value)
            if key in seen:
                continue

            seen.add(key)
            hints.append(f"{label}: {nearby_value}")

    return hints


def _join_non_empty(lines: Iterable[object]) -> str:
    cleaned = [_clean_text(line) for line in lines]
    return "\n".join(line for line in cleaned if line)


def _render_row(values: Iterable[object]) -> str:
    items = [_clean_text(value) for value in values]
    compact = [item for item in items if item]
    return " | ".join(compact)


def _extract_text_from_openpyxl_workbook(file_path: str) -> str:
    workbook = load_workbook(filename=file_path, data_only=True, read_only=True)
    lines: list[str] = []

    try:
        for sheet in workbook.worksheets:
            lines.append(f"Hoja: {sheet.title}")
            sheet_rows = [list(row) for row in sheet.iter_rows(values_only=True)]
            lines.extend(_extract_spreadsheet_field_hints(sheet_rows))
            for row in sheet_rows:
                rendered = _render_row(row)
                if rendered:
                    lines.append(rendered)
    finally:
        workbook.close()

    return _join_non_empty(lines)


def _extract_text_from_tabular_workbook(file_path: str, engine: str) -> str:
    sheets = pd.read_excel(
        file_path,
        sheet_name=None,
        header=None,
        engine=engine,
    )

    lines: list[str] = []
    for sheet_name, frame in sheets.items():
        lines.append(f"Hoja: {sheet_name}")
        rows = [list(row) for row in frame.itertuples(index=False, name=None)]
        lines.extend(_extract_spreadsheet_field_hints(rows))
        for row in rows:
            rendered = _render_row(row)
            if rendered:
                lines.append(rendered)

    return _join_non_empty(lines)


def _extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - guarded by requirements
        raise RuntimeError("python-docx is required to process .docx files") from exc

    document = DocxDocument(file_path)
    lines: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            lines.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            rendered = _render_row(cell.text for cell in row.cells)
            if rendered:
                lines.append(rendered)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            if paragraph.text:
                lines.append(paragraph.text)
        for paragraph in section.footer.paragraphs:
            if paragraph.text:
                lines.append(paragraph.text)

    return _join_non_empty(lines)


def _extract_text_from_doc_via_word(file_path: str) -> str:
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:  # pragma: no cover - guarded by requirements
        raise RuntimeError(
            "pywin32 is required to process .doc files on Windows"
        ) from exc

    path = str(Path(file_path).resolve())

    with _DOC_COM_LOCK:
        pythoncom.CoInitialize()
        word = None
        document = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            document = word.Documents.Open(
                path,
                ConfirmConversions=False,
                ReadOnly=True,
                AddToRecentFiles=False,
                NoEncodingDialog=True,
            )
            return _join_non_empty([document.Content.Text])
        finally:
            if document is not None:
                try:
                    document.Close(False)
                except Exception:
                    logger.debug("Could not close Word document cleanly: %s", file_path)
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    logger.debug("Could not quit Word cleanly for: %s", file_path)
            pythoncom.CoUninitialize()


def _extract_text_from_msg(file_path: str) -> str:
    try:
        import extract_msg
    except ImportError as exc:  # pragma: no cover - guarded by requirements
        raise RuntimeError("extract-msg is required to process .msg files") from exc

    message = extract_msg.Message(file_path)
    try:
        lines = [
            f"Asunto: {message.subject}" if getattr(message, "subject", None) else None,
            f"Remitente: {message.sender}" if getattr(message, "sender", None) else None,
            f"Para: {message.to}" if getattr(message, "to", None) else None,
            f"CC: {message.cc}" if getattr(message, "cc", None) else None,
            f"Fecha: {message.date}" if getattr(message, "date", None) else None,
            getattr(message, "body", None),
        ]
        return _join_non_empty(lines)
    finally:
        close = getattr(message, "close", None)
        if callable(close):
            close()


def extract_text_from_structured_document(file_path: str) -> str:
    logger.info("Starting structured extraction on file: %s", file_path)

    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    source_type = path.suffix.lower().lstrip(".")
    if source_type not in STRUCTURED_TEXT_SOURCE_TYPES:
        raise ValueError(f"Unsupported structured source type: {source_type}")

    try:
        if source_type in SPREADSHEET_SOURCE_TYPES:
            if source_type in {"xlsx", "xlsm"}:
                text = _extract_text_from_openpyxl_workbook(str(path))
            elif source_type == "xls":
                text = _extract_text_from_tabular_workbook(str(path), engine="xlrd")
            else:
                text = _extract_text_from_tabular_workbook(str(path), engine="pyxlsb")
        elif source_type in WORD_SOURCE_TYPES:
            if source_type == "docx":
                text = _extract_text_from_docx(str(path))
            else:
                text = _extract_text_from_doc_via_word(str(path))
        elif source_type in MESSAGE_SOURCE_TYPES:
            text = _extract_text_from_msg(str(path))
        else:  # pragma: no cover - guarded by the source type check above
            text = ""
    except Exception as exc:
        logger.warning("Structured extraction failed for %s: %s", file_path, exc)
        return ""

    result = (text or "").strip()
    logger.info(
        "Structured extraction completed for %s - extracted %d chars",
        file_path,
        len(result),
    )
    return result
